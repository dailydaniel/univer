#!/usr/bin/python
# -*- coding: UTF-8 -*-
#code taken from "http://qaru.site/questions/7036776/how-to-create-a-dataframe-from-a-table-in-a-word-document-docx-file-using-pandas"

import sys
import argparse

def createParser():
    parser = argparse.ArgumentParser()
    parser.add_argument('-d', '--data', default=0, type=int)
    return parser

import pandas as pd
import io
import csv
from docx import Document
import re
from difflib import ndiff
from sklearn.metrics.pairwise import cosine_similarity


####################################################################################################


def read_docx_tables(filename, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to 'pd.read_csv()' function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise


####################################################################################################


def delta(n, df, write = False):
    raw = (df[df.columns[3]][n], df[df.columns[4]][n], df[df.columns[5]][n])

    if write:
        print(raw, '\n')

    step1 = ndiff(raw[0], raw[1])
    dlt0 = []
    dlt1 = []
    if write:
        print('{0} -> {1}\n'.format(dlt[0], dlt[1]))

    for el in step1:
        if el[0] == '-':
            dlt0.append(el[2:])
        if el[0] == ' ' and dlt0:
            dlt0.append(el[2:])

        if el[0] == '+':
            dlt1.append(el[2:])
        if el[0] == ' ' and dlt1:
            dlt1.append(el[2:])

        if write:
            print(el)

    if write:
        print('\n', dlt0, '\n', dlt1, '\n')

    to_delete = []

    step2 = ndiff(raw[1], dlt1)

    for el in step2:
        if el[0] == '-':
            to_delete.append(el[2:])
        else:
            break

        if write:
            print(el)

    if write:
        print()

    t = []

    for i, el in enumerate(raw[2]):
        if i < len(to_delete) and el == to_delete[i]:
            t.append('')
        else:
            t.append(el)

    dlt2 = [x for x in t if x]
    if write:
        print(dlt2)

    return dlt0, dlt1, dlt2


####################################################################################################


def ha(source):
    pattern = r'[\(\)\[\]]'
    res = re.sub(pattern, '', str(source)).split(', ')
    return [re.sub(r"[']", '', x) for x in res]


####################################################################################################


def listReader(raw: str, write = False) -> list:
    if raw[0] == '[':
        raw = re.sub("[\[\]'\) ]", '', raw) #re.sub("[\[\]' ]", '', raw)
        return raw.split(',')

    if write:
        print(raw)
    raw = raw[1:].split('], ')
    return [listReader(x) for x in raw]


####################################################################################################

def affinity(f1: str, f2: str) -> float:
    aff = 0.0
    gl = ['A', 'E', 'I', 'Y', 'O', 'U']
    sogl = ['Z', 'R', 'J', 'G', 'B', 'V', 'W', 'D', 'L', 'M', 'N', 'C', 'H', 'T', 'K', 'P', 'F', 'S']
    zv = ['Z', 'R', 'J', 'G', 'V', 'W', 'D']
    ntr = ['N', 'M', 'L', 'C']
    glh = ['H', 'T', 'K', 'P', 'F', 'S']

    if f1[0] == f2[0] or f1[0] in ['V', 'W'] and f2[0] in ['V', 'W'] or f1[0] in ['J', 'G'] and f2[0] in ['J', 'G']:
        aff += 0.9
    elif f1[0] in gl and f2[0] in gl or f1[0] in sogl and f2[0] in sogl:
        aff += 0.2
        if f1[0] in gl and f2[0] in gl:
            aff += 0.25
        if f1[0] in zv and f2[0] in zv or f1[0] in ntr and f2[0] in ntr or f1[0] in glh and f2[0] in glh:
            aff += 0.25

    f1_a = None
    if f1[-1] in ['0', '1', '2']:
        f1_a = f1[-1]
        f1 = f1[:-1]

    f2_a = None
    if f2[-1] in ['0', '1', '2']:
        f2_a = f2[-1]
        f2 = f2[:-1]

    accent = False
    if f1_a and f2_a:
        if f1_a == f2_a:
            accent = True

    aff += 0.05 if accent else 0.0

    overlap = 0
    mb = []
    if len(f1) == len(f2) == 2:
        mb = [(1, 1)]#[(1, 1), (0, 1), (1, 0)]
    elif len(f1) == 2 and len(f2) == 1:
        mb = [(1, 0)]
        if aff == 0.9: #
            aff += 0.05 #
    elif len(f1) == 1 and len(f2) == 2:
        mb = [(0, 1)]
        if aff == 0.9: #
            aff += 0.05 #
    elif len(f1) == len(f2) == 1:
        if f1[0] == f2[0]:
            aff += 0.1

    if mb:
        for i, j in mb:
            if f1[i] == f2[j]:
                overlap += 1

        aff += overlap / (len(mb) * 2 * 10)

    return round(aff, 3)


####################################################################################################


def cossim(arr1: list, arr2: list, rnd = 5) -> float:
    return round(cosine_similarity([arr1], [arr2])[0][0], rnd)


####################################################################################################


if __name__ == '__main__':
    parser = createParser()
    namespace = parser.parse_args(sys.argv[1:])
    print(namespace.data)
