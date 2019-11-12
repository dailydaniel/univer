
# coding: utf-8

# In[1]:


import math
import numpy

import matplotlib.mlab as mlab
import matplotlib.pyplot as plt

import scipy
import scipy.stats as pystats
from scipy.stats import norm as pynorm
from scipy.stats import uniform as pyuni

from docx import Document


# In[2]:


def strm(a):
    n = len(a)
    b = a[0]
    m = len(b)
    for i in range(n):
        for k in range(m):
            if type(a[i][k]) == int:
                a[i][k] = str(a[i][k])
            elif type(a[i][k]) != str:
                a[i][k] = str('%.5f'% a[i][k])


# In[3]:


def onetablein(docname, n):
    document = Document(docname)
    table = document.tables[n]
    array = []
    for i in range(3,len(table.rows)):
        row = table.rows[i]
        for cell in row.cells:
            array.append(float(cell.text.replace(',','.')))
    return array


# In[4]:


krit_tbl = {
    '4' : 9.5,
    '5' : 11.1,
    '6' : 12.6,
    '7' : 14.1,
    '8' : 15.5}


# In[5]:


ln = onetablein('Data/tables_9.docx', 0)

lu = onetablein('Data/tables_9.docx', 1)

alfa = 0.05
a = 0.1
b = 6.1


# In[6]:


class stat(object):
    def __init__(self, array, m, ao = None, am = None):
        self.m = m
        if ao == None:
            self.ao = min(array)
        else:
            self.ao = ao
        if am == None:
            self.am = max(array)
        else:
            self.am = am
        self.h = (self.am - self.ao)/self.m
        self.num = [self.ao + (i + 0.5)*self.h  for i in range(m)]
        self.interval = [self.ao + i*self.h for i in range(m+1)]
        self.cnt = [0 for i in range(m)]
        for i in array:
            for k in range(m):
                if i <= self.interval[k+1]:
                    break
            self.cnt[k] = self.cnt[k] + 1
        self.war = [i/len(array) for i in self.cnt]


# In[7]:


class NormSample(object):
    def __init__(self, array):
        self.array = [i for i in array]
        self.n = len(array)
        self.m = 1 + int(math.log2(self.n))
        self.stats = stat(array, self.m)
    
    def mean(self):
        s = 0
        for i in range(self.m):
            s = s + self.stats.war[i]*self.stats.num[i]
        return s
    
    def disp(self):
        s = 0
        for i in range(self.m):
            s = s + self.stats.war[i]*(self.stats.num[i]**2)
        return s - (self.stats.h**2)/12 - self.mean()**2
    
    def devi(self):
        return self.disp()**0.5
    
    def graf(self, hist = True, theory = True, filename = 'Data/graf_norm.png'):
        dx = 0.005
        xmin = self.stats.ao - self.stats.h
        xmax = self.stats.am + self.stats.h
        if hist:
            plt.bar(
                [i for i in self.stats.num],
                [i/self.stats.h for i in self.stats.war],
                color = 'blue', edgecolor = 'black', 
                linewidth = 1, alpha = 0.75, width = self.stats.h
            )
        if theory:
            xlist = numpy.arange(xmin, xmax, dx) # mlab.frange(xmin, xmax, dx)
            ylist = [pystats.norm.pdf(xz, loc = self.mean(), scale = self.devi()) for xz in xlist]
            plt.plot(xlist, ylist, color = 'red', linewidth = 2)
        plt.savefig(filename)
        plt.clf()
    
    def out_1(self):
        l = []
        for k in range(len(self.stats.interval)):
            q = []
            ak = self.stats.interval[k]
            q.append(k)
            q.append(ak)
            q.append((ak - self.mean())/self.devi())
            q.append(pynorm.pdf(q[2])/self.devi())
            q.append(pynorm.cdf(q[2]))
            if k == 0:
                q.append('-')
            elif k == 1:
                q.append(q[4])
            elif k == self.m:
                q.append(1 - l[k-1][4])
            else:
                q.append(q[4] - l[k-1][4])
            l.append(q)
        return l
    
    def out_2(self):
        out_1 = self.out_1()
        l = []
        f4 = 0
        f5 = 0
        for k in range(self.m):
            q = []
            q.append(k+1)
            s = ''
            if k == 0:
                s = s + '['
            else:
                s = s + '('
            s = s + str('%.5f'% out_1[k][1]) + '; ' + str('%.5f'% out_1[k+1][1]) + ']'
            q.append(s)
            q.append(self.stats.war[k])
            q.append(out_1[k+1][5])
            q.append(abs(q[2]-q[3]))
            q.append(self.n*(q[4]**2)/q[3])
            if q[4] > f4:
                f4 = q[4]
            f5 = f5 + q[5]
            l.append(q)
        l.append(['','','','',f4,f5])
        return l
        
        
    def hi2(self):
        s = 0
        n = [k for k in self.stats.cnt]
        p = [k[3] for k in self.out_2()]
        for k in range(self.m):
            s = s + ((n[k]-self.n*p[k])**2)/(self.n*p[k])
        return s
    
    def krit(self, tbl):
        l = self.m - 3
        if self.hi2() > tbl[str(l)]:
            return False
        else:
            return True


# In[8]:


sn = NormSample(ln)
print(sn.mean())
print(sn.disp())
print(sn.devi(), '\n')
for i in sn.out_1():
    print(i)
print()
for i in sn.out_2():
    print(i)
print()
print(sn.hi2())
print(sn.krit(krit_tbl))
sn.graf()


# In[9]:


class UniSample(object):
    def __init__(self, array, a, b):
        self.array = [i for i in array]
        self.n = len(array)
        self.a = a
        self.b = b
        self.m = 1 + int(math.log2(self.n))
        self.stats = stat(array, self.m, ao = a, am = b)
    
    
    def mean(self):
        s = 0
        for i in range(self.m):
            s = s + self.stats.war[i]*self.stats.num[i]
        return s
    
    def disp(self):
        s = 0
        for i in range(self.m):
            s = s + self.stats.war[i]*(self.stats.num[i]**2)
        return s - (self.stats.h**2)/12 - self.mean()**2
    
    def devi(self):
        return self.disp()**0.5
    
    def graf(self, hist = True, theory = True, filename = 'Data/graf_uni.png'):
        dx = 0.005
        xmin = self.stats.ao - self.stats.h
        xmax = self.stats.am + self.stats.h
        
        if hist:
            plt.bar(
                [i for i in self.stats.num],
                [i/self.stats.h for i in self.stats.war],
                color = 'blue', edgecolor = 'black', 
                linewidth = 1, alpha = 0.75, width = self.stats.h
            )
        if theory:
            xlist = numpy.arange(xmin, xmax, dx) # mlab.frange(xmin, xmax, dx)
            ylist = [pystats.uniform.pdf(xz, loc = self.a, scale = self.b - self.a) for xz in xlist]
            plt.plot(xlist, ylist, color = 'red', linewidth = 2)
        plt.savefig(filename)
        plt.clf()
        
    def out_3(self):
        l = []
        f4 = 0
        f5 = 0
        for k in range(self.m):
            q = []
            q.append(k+1)
            s = ''
            if k == 0:
                s = s + '['
            else:
                s = s + '('
            s = s + str('%.5f'% self.stats.interval[k]) + '; ' + str('%.5f'% self.stats.interval[k+1]) + ']'
            q.append(s)
            q.append(self.stats.war[k])
            q.append(1/self.m)
            q.append(abs(q[2]-q[3]))
            q.append(self.n*(q[4]**2)/q[3])
            l.append(q)
            if q[4] > f4:
                f4 = q[4]
            f5 = f5 + q[5]
        l.append(['','','','',f4,f5])
        return l
    
    def hi2(self):
        s = 0
        n = [k for k in self.stats.cnt]
        p = [k[3] for k in self.out_3()]
        for k in range(self.m):
            s = s + ((n[k]-self.n*p[k])**2)/(self.n*p[k])
        return s
    
    def krit(self, tbl):
        l = self.m - 3
        if self.hi2() > tbl[str(l)]:
            return False
        else:
            return True


# In[10]:


su = UniSample(lu,a,b)
print(su.mean())
print(su.disp())
print(su.devi())
print()
for k in su.out_3():
    print(k)
print()
print(su.hi2())
print(su.krit(krit_tbl))
su.graf()


# In[11]:


head_1_str = [
    'k',
    'ak',
    '(ak-a)/s',
    '1/s*f((ak-a)/s)',
    'F((ak-a)/s)',
    'pk'
]
head_2_str = [
    'k',
    'Интервал',
    'wk',
    'pk',
    '|wk-pk|',
    'N(wk-pk)^2/pk'
]


# In[12]:


def tabler(document, out, head = None):
    tl = []
    if head != None:
        tl.append(head)
    for i in out:
        tl.append(i)
    strm(tl)
    
    table = document.add_table(rows = len(tl),cols = len(tl[0]))
    for i in range(len(tl)):
        hdr_cells = table.rows[i].cells
        for k in range(len(tl[0])):
            if type(tl[i][k]) == str:
                hdr_cells[k].text = tl[i][k]
            else:
                pass #LaTeh
    


# In[13]:


def doc_task_1(document, sn):
    document.add_paragraph('Задание 1)')
    
    document.add_paragraph('Полученная выборка:')
    l = [i for i in sn.array]
    tbl = []
    q = []
    for i in range(len(l)):
        q.append(l[i])
        if (len(q) == 10) or (i+1 == len(l)):
            while(len(q) != 10):
                q.append('')
            tbl.append(q)
            q = []
    tabler(document, tbl)
    
    document.add_paragraph('')
    document.add_paragraph('Упорядоченная выборка:')
    l.sort()
    tbl = []
    q = []
    for i in range(len(l)):
        q.append(l[i])
        if (len(q) == 10) or (i+1 == len(l)):
            while(len(q) != 10):
                q.append('')
            tbl.append(q)
            q = []
    tabler(document, tbl)
    
    document.add_paragraph('')
    document.add_paragraph('Группированная выборка (интервальный вариационный ряд):')
    tbl = [['(ai-1,ai]'],['ni'],['wi']]
    tbl[0].extend([i[1] for i in sn.out_2() if i[1] != ''])
    tbl[1].extend(sn.stats.cnt)
    tbl[2].extend(sn.stats.war)
    tabler(document, tbl)
    document.add_paragraph('<math>\sum_{i=0}^\{N}\w_i</math> = ' + str('%.5f'% sum(sn.stats.war)))
    
    document.add_paragraph('Математическое ожидание: \tilde{a} = ' + str('%.5f'% sn.mean()))
    document.add_paragraph('Дисперсия: \tilde{\sigma}^2 = '+ str('%.5f'% sn.disp()))
    document.add_paragraph('Среднеквадратическое отклонение: \tilde{\sigma} = ' + str('%.5f'% sn.devi()))
    document.add_paragraph('')
    
    
    tabler(document, sn.out_1(), head_1_str)    
    document.add_paragraph('График плотности нормального распределения, наложенный на гистограмму относительных частот:')
    sn.graf(filename = 'tsk1_graf.png')
    document.add_picture('tsk1_graf.png')
    
    tabler(document, sn.out_2(), head_2_str)
    document.add_paragraph('')
    document.add_paragraph('\chi_B^2 = '+ str('%.5f'% sn.hi2()))
    document.add_paragraph('')


# In[14]:


def doc_task_2(document, sn):
    document.add_paragraph('Задание 2)')
    document.add_paragraph('')
    document.add_paragraph('a = ' + str(sn.a))
    document.add_paragraph('b = ' + str(sn.b))
    document.add_paragraph('')

    
    document.add_paragraph('Полученная выборка:')
    l = [i for i in sn.array]
    tbl = []
    q = []
    for i in range(len(l)):
        q.append(l[i])
        if (len(q) == 10) or (i+1 == len(l)):
            while(len(q) != 10):
                q.append('')
            tbl.append(q)
            q = []
    tabler(document, tbl)
    
    document.add_paragraph('')
    document.add_paragraph('Упорядоченная выборка:')
    l.sort()
    tbl = []
    q = []
    for i in range(len(l)):
        q.append(l[i])
        if (len(q) == 10) or (i+1 == len(l)):
            while(len(q) != 10):
                q.append('')
            tbl.append(q)
            q = []
    tabler(document, tbl)
    
    document.add_paragraph('')
    document.add_paragraph('Группированная выборка (интервальный вариационный ряд):')
    tbl = [['(ai-1,ai]'],['ni'],['wi']]
    tbl[0].extend([i[1] for i in sn.out_3() if i[1] != ''])
    tbl[1].extend(sn.stats.cnt)
    tbl[2].extend(sn.stats.war)
    tabler(document, tbl)
    document.add_paragraph('<math>\sum_{i=0}^\{N}\w_i</math> = ' + str('%.5f'% sum(sn.stats.war)))
    
    document.add_paragraph('Математическое ожидание: \tilde{a} = ' + str('%.5f'% sn.mean()))
    document.add_paragraph('Дисперсия: \tilde{\sigma}^2 = '+ str('%.5f'% sn.disp()))
    document.add_paragraph('Среднеквадратическое отклонение: \tilde{\sigma} = ' + str('%.5f'% sn.devi()))
    document.add_paragraph('')
    
    document.add_paragraph('График плотности равномерного распределения, наложенный на гистограмму относительных частот:')
    sn.graf(filename = 'tsk2_graf.png')
    document.add_picture('tsk2_graf.png')
    
    
    tabler(document, sn.out_3(), head_2_str)
    document.add_paragraph('')
    document.add_paragraph('\chi_B^2 = '+ str('%.5f'% sn.hi2()))
    document.add_paragraph('')


# In[15]:


def doc_anal(document, sn, su, krit_table):
    NormResult = {
        True : 'Гипотеза о соответствии выборки нормальному распределению не противоречит экспериментальным данным (т.е. может быть принята) при уровне значимости alpha = 0,05. ',
        False : 'Гипотеза о соответствии выборки нормальному распределению противоречит экспериментальным данным (т.е. не может быть принята) при уровне значимости alpha = 0,05. '
    }
    UniResult = {
        True: 'Гипотеза о соответствии выборки равномерному распределению на отрезке [a, b] не противоречит экспериментальным данным (т.е. может быть принята) при уровне значимости alpha = 0,05. ',
        False: 'Гипотеза о соответствии выборки равномерному распределению на отрезке [a, b] противоречит экспериментальным данным (т.е. не может быть принята) при уровне значимости alpha = 0,05. '
    }
    
    document.add_paragraph('Анализ результатов и выводы')    
    document.add_paragraph('')
    document.add_paragraph('Таблица критических значений:')# \chi_{кр,\alpha}^2 (l) 
    tbl = [['l'],['chi']]#'\chi_{кр,\alpha}^2 (l)'
    tbl[0].extend([i for i in krit_table])
    tbl[1].extend([str(krit_table[i]) for i in krit_table])
    tabler(document,tbl)
    document.add_paragraph('')    
    document.add_paragraph('1) Нормальное распределение')
    document.add_paragraph('\chi_B^2 = ' + str('%.5f'% sn.hi2()))
    document.add_paragraph('chi(' + str(sn.m-3) + ') = ' + str(krit_table[str(sn.m-3)]) )#'\chi_{кр,\alpha}^2 (l)'
    document.add_paragraph(NormResult[sn.krit(krit_table)])    

    document.add_paragraph('')    
    document.add_paragraph('2) Равномерное распределение')
    document.add_paragraph('\chi_B^2 = ' + str('%.5f'% su.hi2()))
    document.add_paragraph('chi(' + str(su.m-3) + ') = ' + str(krit_table[str(su.m-3)]) )#'\chi_{кр,\alpha}^2 (l)'
    document.add_paragraph(NormResult[su.krit(krit_table)])    


# In[16]:


def doccreator(sn, su, krit_tbl, filename = 'final_doc.docx'):
    document = Document()
    doc_task_1(document, sn)
    document.add_paragraph('')
    doc_task_2(document, su)
    document.add_paragraph('')
    doc_anal(document, sn, su, krit_tbl)
    document.save(filename)
    print('Done!')


# In[17]:


doccreator(sn, su, krit_tbl, filename = 'Data/final_doc.docx')


# In[ ]:




